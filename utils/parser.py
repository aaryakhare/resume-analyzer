"""
parser.py
Extracts raw text from an uploaded resume file (PDF, DOCX, or TXT).
"""
import io
from PyPDF2 import PdfReader
from docx import Document


class UnsupportedFileType(Exception):
    pass


def extract_text(file_storage, filename: str) -> str:
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    raw = file_storage.read()

    if ext == "pdf":
        return _extract_pdf(raw)
    elif ext == "docx":
        return _extract_docx(raw)
    elif ext == "txt":
        return raw.decode("utf-8", errors="ignore")
    else:
        raise UnsupportedFileType(f"Unsupported file type: .{ext}")


def _extract_pdf(raw_bytes: bytes) -> str:
    reader = PdfReader(io.BytesIO(raw_bytes))
    pages_text = []
    for page in reader.pages:
        text = page.extract_text() or ""
        pages_text.append(text)
    return "\n".join(pages_text)


def _extract_docx(raw_bytes: bytes) -> str:
    doc = Document(io.BytesIO(raw_bytes))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)